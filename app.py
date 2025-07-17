from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def generate_invoice_word(data, output_path):
    doc = Document()
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    # 标题
    title = doc.add_heading('Commercial Invoice', level=0)
    title.alignment = 1  # 居中
    subtitle = doc.add_paragraph('商业发票', style='Normal')
    subtitle.alignment = 1  # 居中

    # 表格
    table = doc.add_table(rows=12, cols=7)
    table.style = 'Table Grid'

    # 第一行
    table.cell(0, 0).merge(table.cell(0, 2))
    table.cell(0, 0).text = 'Waybill NO\n运单号: {}'.format(data['waybill_no'])
    table.cell(0, 3).merge(table.cell(0, 6))
    table.cell(0, 3).text = 'Date of Exportation\n出口日期: {}'.format(data['date_of_exportation'])
    

    # # 第二行
    table.cell(1, 0).merge(table.cell(1, 2))
    table.cell(1, 0).text = 'Export company\n发件公司名:\n{}'.format(data['company_name'])
    table.cell(1, 3).merge(table.cell(1, 6))
    table.cell(1, 3).text = 'Contact person\n发件联系人: LUCY YU'


    # # 第三行
    table.cell(2, 0).merge(table.cell(2, 2))
    table.cell(2, 0).text = 'Export Address:\n发件地址:\nFloor 1, Building 7, Alley 13, Luozuxia Village, Luozu Community, Baoan District, Shenzhen, China'
    table.cell(2, 3).merge(table.cell(2, 6))
    table.cell(2, 3).text = 'Phone, Fax or E-mail\n电话: 0086 13411260416'

    # # 第四行
    table.cell(3, 0).merge(table.cell(3, 2))
    table.cell(3, 0).text = 'Importer company\n收件公司名:'
    table.cell(3, 3).merge(table.cell(3, 6))
    table.cell(3, 3).text = 'Contact person\n收件联系人: {}'.format(data['import_contact'])
   

    # # 第五行
    table.cell(4, 0).merge(table.cell(4, 6))
    table.cell(4, 0).text ='Import Address:\n收件地址:\n{}'.format(data['import_address'])

    # # 第六行
    table.cell(5, 0).merge(table.cell(5, 2))
    table.cell(5, 0).text = 'Phone, Fax or E-mail:\n电话: {}'.format(data['import_phone'])
    table.cell(5, 3).merge(table.cell(5, 6))
    table.cell(5, 3).text = 'Zip code\n邮编: {}'.format(data['import_zip'])

     # # 第七行
    table.cell(6, 0).merge(table.cell(6, 1))
    table.cell(6, 0).text = 'Country of Origin\n原产国: China'
    table.cell(6, 2).merge(table.cell(6, 4))
    table.cell(6, 2).text = 'Country of Export\n出口国: China'
    table.cell(6, 5).merge(table.cell(6, 6))
    table.cell(6, 5).text = 'Country of Import\n进口国: USA'
    


    # # 第七行（表头）
    headers = [
        'No. of Pack\n件数',
        'Descriptions\n货物描述\n(HS CODE: 3923210000)',
        'Materials\n材料成份',
        'Application\n产品用途（包装）',
        'Qty\n数量',
        'Unit Price\n单价(USD)',
        'Total Price\n总价(USD)'
    ]
    for i, h in enumerate(headers):
        table.cell(7, i).text = h

    # 动态插入商品数据，至少保留两行
    items = data.get('items', [])
    min_rows = 3
    item_rows = max(len(items), min_rows)
    for idx in range(item_rows):
        row = 8 + idx
        if idx < len(items):
            item = items[idx]
            table.cell(row, 0).text = str(item.get('no_of_pack', ''))
            table.cell(row, 1).text = item.get('description', '')
            table.cell(row, 2).text = item.get('material', '')
            table.cell(row, 3).text = item.get('application', '')
            table.cell(row, 4).text = str(item.get('qty', ''))
            table.cell(row, 5).text = str(item.get('unit_price', ''))
            table.cell(row, 6).text = str(item.get('total_price', ''))
        else:
            for col in range(7):
                table.cell(row, col).text = ''

    # 合计行，始终在数据行后
    sum_row = 8 + item_rows
    table.cell(sum_row, 0).merge(table.cell(sum_row, 2))
    table.cell(sum_row, 0).text = 'Total No. of Package\n总件数: {}'.format(data['total_no_of_package'])
    table.cell(sum_row, 3).merge(table.cell(sum_row, 6))
    # 合计金额自动计算
    total_price = sum(float(item.get('total_price', 0) or 0) for item in items)
    table.cell(sum_row, 3).text = 'Total Price {:.2f}\n总价(USD)'.format(total_price)
    # # 备注
    doc.add_paragraph('备注：请用英文填写，材料成份、产品用途必填')

    doc.save(output_path)

# 示例调用
data = {
    'waybill_no': '6632948000',
    'date_of_exportation': '4-16',
    'company_name':'Shenzhen HengMao Pack Co.,Ltd',
    'import_contact': 'Albert Cruz',
    'import_address': '2125 quimby ave Bronx ny 10473',
    'import_phone': '+1 954-559-8137',
    'import_zip': '10473',
    'no_of_pack': 1,
    'qty': '504pcs',
    'total_no_of_package': 1,
    'items': [
        {'no_of_pack': 1, 'description': 'Plasticbottle', 'material': 'PP', 'application': 'packing', 'qty': '504pcs', 'unit_price': '0.34', 'total_price': '171.36'}
    ]
}
generate_invoice_word(data, 'invoice.docx')